import os
import logging
from typing import List, Dict, Any, Optional, TypedDict
from pathlib import Path
import json

# LangChain imports
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_ollama import OllamaLLM
from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document

# LangGraph imports
from langgraph.graph import StateGraph, END

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global embedding cache to avoid reloading
_global_embeddings_cache = None

# ===================== WORKFLOW STATE =====================

class WorkflowState(TypedDict):
    """State for the question generation workflow"""
    file_path: str
    file_type: str  # 'pdf', 'txt', 'doc', 'docx'
    user_prompt: str
    documents: List[Document]
    is_large_document: bool
    vector_store: Optional[Any]
    relevant_content: str
    generated_response: str
    error: Optional[str]
    metadata: Dict[str, Any]

# ===================== DOCUMENT PROCESSOR =====================

class MedicalDocumentProcessor:
    """Handles medical document loading and processing using LangChain"""
    
    def __init__(self):
        # Optimal settings for medical content
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2200,        # Optimal for medical context
            chunk_overlap=350,      # Preserves medical relationships  
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def _load_embeddings(self):
        """Load embeddings only when needed for RAG - with global caching"""
        global _global_embeddings_cache
        
        if _global_embeddings_cache is None:
            logger.info("Loading embedding model for RAG processing...")
            _global_embeddings_cache = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-mpnet-base-v2"
            )
            logger.info("Embedding model loaded successfully")
        else:
            logger.info("Using cached embedding model")
            
        return _global_embeddings_cache
    
    def load_document(self, file_path: str, file_type: str) -> List[Document]:
        """Load document based on file type"""
        if file_type == 'pdf':
            return self._load_pdf(file_path)
        elif file_type == 'txt':
            return self._load_txt(file_path)
        elif file_type in ['doc', 'docx']:
            return self._load_word(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _load_pdf(self, pdf_path: str) -> List[Document]:
        """Load PDF using LangChain's PyMuPDFLoader"""
        loader = PyMuPDFLoader(pdf_path)
        documents = loader.load()
        
        # Clean documents
        for doc in documents:
            doc.page_content = self._clean_text(doc.page_content)
        
        logger.info(f"Loaded {len(documents)} pages from PDF")
        return documents
    
    def _load_txt(self, txt_path: str) -> List[Document]:
        """Load TXT file"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Clean content
            cleaned_content = self._clean_text(content)
            
            # Create document with metadata
            document = Document(
                page_content=cleaned_content,
                metadata={
                    "source": txt_path,
                    "file_type": "txt",
                    "length": len(cleaned_content)
                }
            )
            
            logger.info(f"Loaded TXT file with {len(cleaned_content)} characters")
            return [document]
            
        except Exception as e:
            logger.error(f"Error loading TXT file: {e}")
            raise
    
    def _load_word(self, word_path: str) -> List[Document]:
        """Load DOC/DOCX file"""
        try:
            # Try to import python-docx for Word documents
            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise ImportError("python-docx is required for DOC/DOCX files. Install with: pip install python-docx")
            
            # Load DOCX file
            doc = DocxDocument(word_path)
            
            # Extract text from all paragraphs
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
            
            # Clean content
            cleaned_content = self._clean_text(content)
            
            # Create document with metadata
            document = Document(
                page_content=cleaned_content,
                metadata={
                    "source": word_path,
                    "file_type": "docx",
                    "length": len(cleaned_content),
                    "paragraphs": len(doc.paragraphs)
                }
            )
            
            logger.info(f"Loaded Word document with {len(cleaned_content)} characters")
            return [document]
            
        except Exception as e:
            logger.error(f"Error loading Word document: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text"""
        import re
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)\[\]\/\%°]', '', text)
        return text.strip()
    
    def should_use_rag(self, documents: List[Document]) -> bool:
        """Determine if RAG should be used based on document size"""
        total_content = sum(len(doc.page_content) for doc in documents)
        return total_content > 2000  # ~1.5 pages for medical precision
    
    def create_vector_store(self, documents: List[Document]) -> Chroma:
        """Create vector store from documents - only loads embeddings when needed"""
        # Load embeddings only when we actually need RAG
        embeddings = self._load_embeddings()
        
        # Clean up old vector stores (older than 1 hour)
        self._cleanup_old_vector_stores()
        
        # Use unique directory for each session to avoid file locking issues
        import uuid
        import time
        vector_dir = f"./chroma_pdf/chroma_medical_db_{int(time.time())}_{str(uuid.uuid4())[:8]}"
        
        texts = self.text_splitter.split_documents(documents)
        vector_store = Chroma.from_documents(
            documents=texts,
            embedding=embeddings,
            persist_directory=vector_dir
        )
        logger.info(f"Created fresh vector store with {len(texts)} chunks in {vector_dir}")
        return vector_store
    
    def _cleanup_old_vector_stores(self):
        """Clean up old vector store directories to save disk space"""
        try:
            import glob
            import time
            import shutil
            
            # Find all old vector store directories
            pattern = "./chroma_medical_db_*"
            old_dirs = glob.glob(pattern)
            current_time = time.time()
            
            for dir_path in old_dirs:
                try:
                    # Extract timestamp from directory name
                    timestamp = int(dir_path.split('_')[3])
                    # Delete if older than 1 hour (3600 seconds)
                    if current_time - timestamp > 3600:
                        shutil.rmtree(dir_path)
                        logger.info(f"Cleaned up old vector store: {dir_path}")
                except (ValueError, IndexError, OSError):
                    # Skip if can't parse timestamp or delete fails
                    pass
        except Exception as e:
            # Don't fail if cleanup fails
            logger.warning(f"Failed to cleanup old vector stores: {e}")

# ===================== PROMPT TEMPLATE =====================

class MedicalPromptTemplate:
    """Medical education prompting template"""
    
    @staticmethod
    def get_template() -> ChatPromptTemplate:
        return ChatPromptTemplate.from_messages([
            ("system", """You are a doctor conducting a clinical interview with a patient.

CRITICAL RULES:
- Use ONLY the information provided in the medical content below
- Do NOT use your general medical knowledge
- FIRST check if the requested topic exists in the medical content
- If the content doesn't contain information about the requested topic, you MUST respond with: "The provided document does not contain sufficient information about [topic]"
- Only create questions if the topic is clearly present in the medical content
- Always format questions as doctor asking patient
- Use simple, clear language
- No explanations or correct answers
- Always use numbered options (1, 2, 3) - never A, B, C, D"""),
            
            ("human", """Based STRICTLY on the following medical content, create clinical interview questions for this request:

USER REQUEST: {user_prompt}

MEDICAL CONTENT:
{content}

CRITICAL INSTRUCTIONS: 
- FIRST: Check if the requested topic is mentioned in the medical content above
- If the topic is NOT found in the content, respond ONLY with: "The provided document does not contain sufficient information about [topic]"
- If the topic IS found, then create clinical interview questions using ONLY that content
- Do NOT create questions about different topics than requested
- Do NOT use your general medical knowledge

ALWAYS use this format ONLY if the topic exists in the content:
Question 1: Do you experience [symptom]?
   1. Yes, frequently
   2. Yes, occasionally  
   3. No, never

Question 2: How long have you had [condition]?
   1. Less than 1 week
   2. 1-2 weeks
   3. More than 2 weeks""")
        ])

# ===================== LANGGRAPH WORKFLOW =====================

class MedicalQuestionWorkflow:
    """LangGraph-based workflow for medical content generation"""
    
    def __init__(self):
        self.doc_processor = MedicalDocumentProcessor()
        # Fixed model for consistency
        self.llm = OllamaLLM(model="llama3.1:8b", temperature=0.3,max_tokens=2000)
        self.prompt_template = MedicalPromptTemplate.get_template()
        self.workflow = self._create_workflow()
    
    def _create_workflow(self) -> StateGraph:
        """Create the LangGraph workflow"""
        workflow = StateGraph(WorkflowState)
        
        # Add nodes
        workflow.add_node("load_documents", self._load_documents)
        workflow.add_node("analyze_document_size", self._analyze_document_size)
        workflow.add_node("create_vector_store", self._create_vector_store)
        workflow.add_node("retrieve_content", self._retrieve_content)
        workflow.add_node("use_direct_content", self._use_direct_content)
        workflow.add_node("generate_response", self._generate_response)
        
        # Define workflow edges
        workflow.set_entry_point("load_documents")
        workflow.add_edge("load_documents", "analyze_document_size")
        
        # Conditional routing based on document size
        workflow.add_conditional_edges(
            "analyze_document_size",
            self._should_use_rag,
            {
                "use_rag": "create_vector_store",
                "use_direct": "use_direct_content"
            }
        )
        
        workflow.add_edge("create_vector_store", "retrieve_content")
        workflow.add_edge("retrieve_content", "generate_response")
        workflow.add_edge("use_direct_content", "generate_response")
        workflow.add_edge("generate_response", END)
        
        return workflow.compile()
    
    def _load_documents(self, state: WorkflowState) -> WorkflowState:
        """Load documents"""
        try:
            documents = self.doc_processor.load_document(state["file_path"], state["file_type"])
            state["documents"] = documents
            state["metadata"] = {"total_pages": len(documents), "file_type": state["file_type"]}
        except Exception as e:
            state["error"] = f"Failed to load documents: {e}"
            logger.error(state["error"])
        return state
    
    def _analyze_document_size(self, state: WorkflowState) -> WorkflowState:
        """Analyze if document needs RAG processing"""
        if state.get("error"):
            return state
            
        is_large = self.doc_processor.should_use_rag(state["documents"])
        state["is_large_document"] = is_large
        state["metadata"]["processing_method"] = "RAG" if is_large else "Direct"
        
        # Calculate actual content length for debugging
        total_content = sum(len(doc.page_content) for doc in state["documents"])
        state["metadata"]["total_content_chars"] = total_content
        
        if is_large:
            logger.info(f"Large document detected ({total_content} chars) - will use RAG with embeddings")
        else:
            logger.info(f"Small document detected ({total_content} chars) - using direct processing (no embeddings needed)")
            
        return state
    
    def _should_use_rag(self, state: WorkflowState) -> str:
        """Conditional routing function"""
        return "use_rag" if state["is_large_document"] else "use_direct"
    
    def _create_vector_store(self, state: WorkflowState) -> WorkflowState:
        """Create vector store for large documents"""
        if state.get("error"):
            return state
            
        try:
            vector_store = self.doc_processor.create_vector_store(state["documents"])
            state["vector_store"] = vector_store
            logger.info("Created vector store for RAG processing")
        except Exception as e:
            state["error"] = f"Failed to create vector store: {e}"
        return state
    
    def _retrieve_content(self, state: WorkflowState) -> WorkflowState:
        """Retrieve relevant content using RAG"""
        if state.get("error"):
            return state
            
        try:
            # Get documents with similarity scores to filter irrelevant content
            relevant_docs_with_scores = state["vector_store"].similarity_search_with_score(
                state["user_prompt"], 
                k=7
            )
            
            # Debug: Log similarity scores
            logger.info(f"Similarity scores for query '{state['user_prompt']}':")
            for i, (doc, score) in enumerate(relevant_docs_with_scores):
                logger.info(f"  Chunk {i+1}: Score {score:.3f} - Content preview: {doc.page_content[:100]}...")
            
            # Use similarity threshold to filter truly relevant content
            similarity_threshold = 1.5  # Optimal threshold for medical content
            relevant_docs = [doc for doc, score in relevant_docs_with_scores if score < similarity_threshold]
            
            if not relevant_docs:
                logger.info(f"No chunks found below similarity threshold {similarity_threshold} - topic not in document")
                state["relevant_content"] = "NO_RELEVANT_CONTENT_FOUND"
                state["metadata"]["retrieved_chunks"] = 0
                state["metadata"]["similarity_scores"] = [score for _, score in relevant_docs_with_scores]
            else:
                relevant_content = "\n\n".join([doc.page_content for doc in relevant_docs])
                state["relevant_content"] = relevant_content
                state["metadata"]["retrieved_chunks"] = len(relevant_docs)
                state["metadata"]["similarity_scores"] = [score for _, score in relevant_docs_with_scores[:len(relevant_docs)]]
                logger.info(f"Retrieved {len(relevant_docs)} relevant chunks below similarity threshold")
            
        except Exception as e:
            state["error"] = f"Failed to retrieve content: {e}"
        return state
    
    def _use_direct_content(self, state: WorkflowState) -> WorkflowState:
        """Use direct content for small documents"""
        if state.get("error"):
            return state
            
        full_content = "\n\n".join([doc.page_content for doc in state["documents"]])
        state["relevant_content"] = full_content
        state["metadata"]["content_length"] = len(full_content)
        
        logger.info("Using direct content processing")
        return state
    
    def _generate_response(self, state: WorkflowState) -> WorkflowState:
        """Generate response using LLM"""
        if state.get("error"):
            return state
            
        try:
            # Check if we found relevant content
            if state["relevant_content"] == "NO_RELEVANT_CONTENT_FOUND":
                # Extract topic from user prompt for better error message
                user_prompt = state["user_prompt"].lower()
                topic = "the requested topic"
                if "about" in user_prompt:
                    topic_part = user_prompt.split("about")[-1].strip()
                    if topic_part:
                        topic = topic_part
                
                state["generated_response"] = f"The provided document does not contain sufficient information about {topic}."
                state["metadata"]["status"] = "no_relevant_content"
                logger.info("No relevant content found - returning insufficient information message")
                return state
            
            # Format prompt with relevant content
            formatted_prompt = self.prompt_template.format_messages(
                user_prompt=state["user_prompt"],
                content=state["relevant_content"]
            )
            
            # Generate response
            response = self.llm.invoke(formatted_prompt)
            
            # Extract content based on response type
            if hasattr(response, 'content'):
                state["generated_response"] = response.content
            else:
                state["generated_response"] = str(response)
                
            state["metadata"]["status"] = "success"
            logger.info("Generated response successfully")
            
        except Exception as e:
            state["error"] = f"Failed to generate response: {e}"
        return state
    
    def generate_response(self, file_path: str, file_type: str, user_prompt: str) -> Dict[str, Any]:
        """Main method to generate response"""
        # Initialize state
        initial_state = WorkflowState(
            file_path=file_path,
            file_type=file_type,
            user_prompt=user_prompt,
            documents=[],
            is_large_document=False,
            vector_store=None,
            relevant_content="",
            generated_response="",
            error=None,
            metadata={}
        )
        
        # Run workflow
        final_state = self.workflow.invoke(initial_state)
        
        # Prepare result
        result = {
            "status": "error" if final_state.get("error") else "success",
            "user_prompt": user_prompt,
            "metadata": final_state["metadata"],
            "response": final_state["generated_response"]
        }
        
        if final_state.get("error"):
            result["error"] = final_state["error"]
        
        return result

# ===================== MAIN INTERFACE =====================

class MedicalDocumentQuestionGenerator:
    """Main interface for medical document processing"""
    
    def __init__(self):
        self.workflow = MedicalQuestionWorkflow()
    
    def process(self, file_path: str, prompt: str) -> Dict[str, Any]:
        """Process document with custom prompt"""
        # Detect file type
        file_extension = Path(file_path).suffix.lower().replace('.', '')
        
        # Map file extensions
        if file_extension == 'pdf':
            file_type = 'pdf'
        elif file_extension == 'txt':
            file_type = 'txt'
        elif file_extension in ['doc', 'docx']:
            file_type = file_extension
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        return self.workflow.generate_response(file_path, file_type, prompt)
    
    def save_result(self, result: Dict[str, Any], output_file: str):
        """Save result to file"""
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        logger.info(f"Result saved to {output_file}")

# ===================== CLI INTERFACE (Optional) =====================

def main():
    """Simple CLI interface for testing"""
    import argparse
    
    parser = argparse.ArgumentParser(description="Medical Document Question Generator")
    parser.add_argument("--file", required=True, help="Path to document file (PDF, TXT, DOC, DOCX)")
    parser.add_argument("--prompt", required=True, help="Your custom prompt/request")
    parser.add_argument("--output", help="Output file path")
    
    args = parser.parse_args()
    
    # Initialize processor
    processor = MedicalDocumentQuestionGenerator()
    
    try:
        # Process document with custom prompt
        print(f"Processing file: {args.file}")
        print(f"User prompt: {args.prompt}")
        print(f"Using model: llama3.1:8b")
        
        result = processor.process(args.file, args.prompt)
        
        # Print result
        if result["status"] == "success":
            print("\n" + "="*50)
            print("GENERATED RESPONSE:")
            print("="*50)
            print(result["response"])
            print("="*50)
        else:
            print(f"Error: {result.get('error', 'Unknown error')}")
        
        # Save if output specified
        if args.output:
            processor.save_result(result, args.output)
            print(f"\nResult saved to: {args.output}")
            
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()